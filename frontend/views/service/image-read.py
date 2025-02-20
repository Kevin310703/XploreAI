import re
import streamlit as st
import requests
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image

from config import API_BASE_URL_BACKEND_SERVICE

if not API_BASE_URL_BACKEND_SERVICE:
    raise ValueError("🚨 API base url is not set in the environment variables!")

if "cookie_manager" not in st.session_state:
    import extra_streamlit_components as stx
    st.session_state.cookie_manager = stx.CookieManager()

if "auth" not in st.session_state:
    st.session_state.auth = {
        "logged_in": False,
        "username": None,
        "access_token": None,
        "refresh_token": None
    }

cookie_manager = st.session_state.cookie_manager

if not st.session_state.auth.get("logged_in"):
    st.warning("⚠️ You need to log in to view your profile.")
    st.stop()

access_token = st.session_state.auth.get("access_token")
if not access_token:
    st.error("❌ Access token not found!")
    st.stop()


def process_table_lines(doc, table_lines):
    """
    Xử lý các dòng Markdown biểu diễn bảng và chèn bảng vào tài liệu Word.
    Giả sử bảng có định dạng:
      | header1 | header2 | ... |
      |---------|---------|-----|
      | data1   | data2   | ... |
    """
    if len(table_lines) < 2:
        return

    # Xử lý dòng header
    header_cells = [cell.strip() for cell in table_lines[0].strip().split("|") if cell.strip()]
    num_cols = len(header_cells)
    # Các dòng dữ liệu bắt đầu từ dòng thứ 3 (bỏ qua dòng phân cách thứ 2)
    data_lines = table_lines[2:]
    num_rows = 1 + len(data_lines)  # 1 header + các dòng dữ liệu

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"  # Bạn có thể điều chỉnh style theo ý muốn

    # Điền header
    hdr_cells = table.rows[0].cells
    for idx, cell_text in enumerate(header_cells):
        hdr_cells[idx].text = cell_text

    # Điền dữ liệu các dòng còn lại
    for row_idx, line in enumerate(data_lines, start=1):
        row_cells = [cell.strip() for cell in line.strip().split("|") if cell.strip()]
        cells = table.rows[row_idx].cells
        for col_idx in range(min(len(row_cells), num_cols)):
            cells[col_idx].text = row_cells[col_idx]


def add_markdown_to_doc(doc, markdown_text):
    """
    Xử lý Markdown theo yêu cầu:
      - Nếu dòng bắt đầu bằng dấu '#' sẽ được xử lý thành heading theo mức tương ứng.
      - Nếu dòng có định dạng **Heading:** (chữ được bọc bởi **, có thể kèm dấu ":" cuối dòng) và dòng không có ký tự bullet, chuyển thành heading level 2.
      - Nếu dòng có định dạng bullet (bắt đầu bằng "*" hoặc "-") với nội dung bôi đậm **...:** sẽ được chuyển thành bullet list (với ký hiệu bullet tự hiển thị từ style).
      - Xử lý bảng Markdown và xử lý inline các đoạn bôi đậm.
    """
    lines = markdown_text.splitlines()
    table_buffer = []  # Lưu các dòng thuộc bảng Markdown

    for line in lines:
        stripped = line.strip()

        # Bỏ qua dòng trống
        if not stripped:
            continue

        # Nếu là dòng bảng Markdown (bắt đầu và kết thúc bằng dấu "|")
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(line)
            continue
        else:
            # Nếu có dữ liệu bảng đã thu thập mà dòng hiện tại không thuộc bảng, xử lý bảng
            if table_buffer:
                process_table_lines(doc, table_buffer)
                table_buffer = []

        # --- Xử lý tiêu đề ---
        # 1. Tiêu đề theo Markdown (dòng bắt đầu bằng '#' )
        header_hash_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if header_hash_match:
            level = len(header_hash_match.group(1))
            heading_text = header_hash_match.group(2).strip()
            para = doc.add_heading(heading_text, level=level)
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        # 2. Tiêu đề dạng **Heading:** – nếu dòng không bắt đầu bằng bullet
        if not re.match(r'^[-*]\s+', stripped):
            heading_match = re.match(r'^\*\*(.*?)\*\*\s*:?\s*$', stripped)
            if heading_match:
                heading_text = heading_match.group(1).strip()
                para = doc.add_heading(heading_text, level=2)
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.color.rgb = RGBColor(0, 0, 0)
                continue

        # --- Xử lý bullet list ---
        # Nếu dòng bắt đầu bằng "*" hoặc "-" có định dạng: * **Content:** ... hoặc - **Content:** ...
        bullet_match = re.match(r'^[-*]\s*\*\*(.*?)\*\*\s*:?\s*(.*)$', stripped)
        if bullet_match:
            bullet_text = bullet_match.group(1).strip()
            rest_of_line = bullet_match.group(2)
            para = doc.add_paragraph(style="List Bullet")
            # Không thêm ký tự "• " thủ công vì style List Bullet đã hiển thị bullet
            content = bullet_text
            if not content.endswith(":"):
                content += ":"
            run = para.add_run(content)
            run.bold = True
            if rest_of_line:
                para.add_run(" " + rest_of_line)
            continue

        # Nếu dòng bắt đầu bằng ký tự bullet nhưng không có định dạng bold đặc trưng, loại bỏ ký tự bullet.
        if re.match(r'^[-*]\s+', stripped):
            line = re.sub(r'^[-*]\s+', '', line)

        # --- Xử lý inline bold ---
        para = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                para.add_run(part)

    if table_buffer:
        process_table_lines(doc, table_buffer)


# --- Phần giao diện Streamlit và gọi API ---
with st.form("vqa_form", enter_to_submit=True, clear_on_submit=False, border=False):
    st.title("📷 Image Reading")

    with st.expander("ℹ️ How to Use", expanded=False):
        st.markdown("""
            🔹 **Step 1:** Upload an image by clicking on **"📤 Upload an image"**.  
            🔹 **Step 2:** Type your question related to the image in the text area.  
            🔹 **Step 3:** Click **"Get Description 🔍"** to process your question.  
            🔹 **Step 4:** The AI will analyze the image and provide an answer based on your question.  

            ⚠️ **Note:**  
            - Make sure to upload a clear image for better accuracy.  
            - Provide specific and clear questions to get the best responses.    
            - If the response takes too long, check your internet connection.    
        """)

    uploaded_file = st.file_uploader("📤 Upload an image:", type=["png", "jpg", "jpeg"])
    if uploaded_file:
        image = Image.open(uploaded_file)
        st.image(image, caption="📷 Uploaded Image", use_container_width=True)

    question = st.text_area("❓ Ask a question about the image:", height=150, placeholder="Enter your question")

    submitted = st.form_submit_button("Get Description 🔍")
    if submitted:
        with st.spinner("⏳ Processing..."):
            files = {"image": uploaded_file.getvalue()}
            data = {"question": question}
            headers = {"Authorization": f"Bearer {access_token}"}

            try:
                response = requests.post(
                    f"{API_BASE_URL_BACKEND_SERVICE}/image-text/",
                    headers=headers,
                    files=files,
                    data=data,
                    timeout=120
                )

                if response.status_code == 200:
                    answer = response.json().get("answer", "No answer received.")
                    st.success("✅ Answer:")
                    st.markdown(answer)

                    # Tạo tài liệu Word với cấu trúc định dạng theo yêu cầu
                    doc = Document()
                    
                    # Cập nhật style "Normal"
                    normal_style = doc.styles["Normal"]
                    normal_style.font.name = "Times New Roman"
                    normal_style.font.size = Pt(13)
                    normal_style.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Cập nhật style cho Heading 1 và Heading 2 (có thể điều chỉnh kích cỡ nếu cần)
                    heading1_style = doc.styles["Heading 1"]
                    heading1_style.font.name = "Times New Roman"
                    heading1_style.font.size = Pt(16)
                    heading1_style.font.color.rgb = RGBColor(0, 0, 0)
                    
                    heading2_style = doc.styles["Heading 2"]
                    heading2_style.font.name = "Times New Roman"
                    heading2_style.font.size = Pt(14)
                    heading2_style.font.color.rgb = RGBColor(0, 0, 0)

                    # Xử lý nội dung Markdown trong 'answer'
                    add_markdown_to_doc(doc, answer)

                    # Lưu tài liệu vào bộ nhớ
                    byte_io = io.BytesIO()
                    doc.save(byte_io)
                    byte_io.seek(0)

                    # Cung cấp nút tải file ngoài form
                    st.session_state["download_ready"] = True
                    st.session_state["word_file"] = byte_io
                    st.session_state["word_file_name"] = "image_description.docx"
                else:
                    st.error(f"❌ API error: {response.status_code} - {response.text}")

            except requests.exceptions.Timeout:
                st.error("⏳ API request timed out. Please try again.")

            except requests.exceptions.RequestException as e:
                st.error(f"⚠️ API request failed: {e}")

# Hiển thị nút download nếu file đã sẵn sàng
if st.session_state.get("download_ready", False):
    download_clicked = st.download_button(
        label="📥 Download Answer as Word",
        data=st.session_state["word_file"],
        file_name=st.session_state["word_file_name"],
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    # Nếu người dùng nhấn tải, ẩn nút tải bằng cách cập nhật session_state
    if download_clicked:
        st.session_state["download_ready"] = False
    st.rerun()
    
st.markdown("---")
st.markdown("🚀 **© 2025 X-OR AI GENERATIVE**")
